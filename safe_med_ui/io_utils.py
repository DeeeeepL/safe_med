import json
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

import pandas as pd
from docx import Document


STRUCTURED_EXT = {".csv", ".xlsx", ".xls", ".json"}
SEMI_STRUCTURED_EXT = {".txt", ".docx", ".jsonl"}


@dataclass
class LoadedData:
    kind: str  # "text" | "docx" | "df" | "jsonl"
    path: Path
    df: Optional[pd.DataFrame] = None
    text: Optional[str] = None
    docx_paragraphs: Optional[List[str]] = None
    jsonl_rows: Optional[List[Dict[str, Any]]] = None
    json_obj: Optional[Any] = None


def detect_kind(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt"}:
        return "text"
    if ext in {".docx"}:
        return "docx"
    if ext in {".csv", ".xlsx", ".xls", ".json"}:
        return "df"
    if ext in {".jsonl"}:
        return "jsonl"
    raise ValueError(f"不支持的文件类型: {ext}")


def load_file(path: str) -> LoadedData:
    p = Path(path)
    kind = detect_kind(p)

    if kind == "text":
        return LoadedData(kind="text", path=p, text=p.read_text(encoding="utf-8", errors="ignore"))

    if kind == "docx":
        doc = Document(str(p))
        paras = [para.text for para in doc.paragraphs]
        return LoadedData(kind="docx", path=p, docx_paragraphs=paras)

    if kind == "df":
        ext = p.suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(p, dtype=str, keep_default_na=False)
        elif ext in {".xlsx", ".xls"}:
            df = pd.read_excel(p, dtype=str, keep_default_na=False)
        elif ext == ".json":
            # 更鲁棒地处理 JSON：支持 list[dict], list[scalar], dict[list], dict[scalar]
            text = p.read_text(encoding="utf-8", errors="ignore")
            obj = json.loads(text)
            # 列表情况
            if isinstance(obj, list):
                if len(obj) == 0:
                    df = pd.DataFrame()
                else:
                    first = obj[0]
                    if isinstance(first, dict):
                        # list of dicts -> 每个dict为一行
                        df = pd.DataFrame(obj)
                    else:
                        # list of scalars -> 单列
                        df = pd.DataFrame({"value": obj})
            # 字典情况
            elif isinstance(obj, dict):
                # 如果值都是列表，则视为多列数据
                if all(isinstance(v, list) for v in obj.values()):
                    df = pd.DataFrame(obj)
                else:
                    # 标量字典 -> 单行数据
                    df = pd.DataFrame([obj])
            else:
                # 其他标量类型 -> 单列单行
                df = pd.DataFrame({"value": [obj]})
            # 将所有数据转换为字符串以便脱敏处理
            if not df.empty:
                df = df.astype(str)
        else:
            raise ValueError(f"不支持的结构化类型: {ext}")
        df = df.fillna("")
        # 如果是 JSON 文件，保留原始解析对象以便导出时保持结构
        if ext == ".json":
            return LoadedData(kind="df", path=p, df=df, json_obj=obj)
        return LoadedData(kind="df", path=p, df=df)

    if kind == "jsonl":
        rows: List[Dict[str, Any]] = []
        with p.open("r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                rows.append(json.loads(line))
        return LoadedData(kind="jsonl", path=p, jsonl_rows=rows)

    raise ValueError("未知 kind")


def suggest_output_path(input_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = input_path.stem
    ext = input_path.suffix.lower()
    return out_dir / f"{stem}_deid{ext}"


def save_text(out_path: Path, text: str) -> None:
    out_path.write_text(text, encoding="utf-8")


def save_docx(out_path: Path, paragraphs: List[str]) -> None:
    doc = Document()
    for t in paragraphs:
        doc.add_paragraph(t)
    doc.save(str(out_path))


def save_df(out_path: Path, df: pd.DataFrame) -> None:
    ext = out_path.suffix.lower()
    if ext == ".csv":
        df.to_csv(out_path, index=False, encoding="utf-8-sig")
    elif ext in {".xlsx", ".xls"}:
        df.to_excel(out_path, index=False)
    elif ext == ".json":
        df.to_json(out_path, orient="records", force_ascii=False, indent=2)
    else:
        raise ValueError(f"不支持写出类型: {ext}")


def save_jsonl(out_path: Path, rows: List[Dict[str, Any]]) -> None:
    with out_path.open("w", encoding="utf-8") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")


def save_json(out_path: Path, obj: Any) -> None:
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def get_text_columns(df: pd.DataFrame) -> List[str]:
    cols = []
    for c in df.columns:
        # 只要能转成 str 就按可脱敏文本处理
        cols.append(str(c))
    return cols


def scan_text_files(folder_path: Path) -> List[Path]:
    """
    扫描文件夹，返回所有文本类文件的路径列表
    支持的文本格式：.txt, .docx, .csv, .xlsx, .json, .jsonl
    """
    TEXT_EXTS = {".txt", ".docx", ".csv", ".xlsx", ".xls", ".json", ".jsonl"}
    text_files = []
    
    for file_path in folder_path.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in TEXT_EXTS:
            text_files.append(file_path)
    
    # 按路径排序便于显示
    return sorted(text_files)


def get_relative_path(file_path: Path, base_path: Path) -> str:
    """获取相对于基础路径的相对路径"""
    try:
        return str(file_path.relative_to(base_path))
    except ValueError:
        return str(file_path)
